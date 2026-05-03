import re
import sys
from docx import Document
from docx.shared import Pt

def markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            parse_bold_inline(p, line[2:])
        elif line[0].isdigit() and line[1:3] == '. ':
            p = doc.add_paragraph(style='List Number')
            parse_bold_inline(p, line[3:])
        else:
            p = doc.add_paragraph()
            parse_bold_inline(p, line)
            
    doc.save(docx_path)

def parse_bold_inline(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
        else:
            paragraph.add_run(part)

if __name__ == "__main__":
    markdown_to_docx(r"e:\export_feature_documentation.md", r"e:\export_feature_documentation.docx")
